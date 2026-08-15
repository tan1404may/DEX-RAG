from pathlib import Path
from typing import Union, BinaryIO
from io import BytesIO

from docx import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .base import BaseExtractor
from ..models import ExtractedDocument, ExtractedSection, DocumentType


class DOCXExtractor(BaseExtractor):
    """Extracts text and structure from .docx files."""
    
    supported_types = [DocumentType.DOCX]
    
    def extract(self, source: Union[str, Path, BinaryIO], **kwargs) -> ExtractedDocument:
        if isinstance(source, (str, Path)):
            doc = DocxDocument(str(source))
            file_size = Path(source).stat().st_size
            mime = self._detect_mime(source)
            source_name = str(source)
        else:
            source.seek(0)
            doc = DocxDocument(BytesIO(source.read()))
            source.seek(0, 2)
            file_size = source.tell()
            source.seek(0)
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            source_name = kwargs.get("filename", "uploaded.docx")
        
        return self._process_docx(doc, source_name, file_size, mime)
    
    def extract_from_bytes(self, data: bytes, filename: str = None, **kwargs) -> ExtractedDocument:
        doc = DocxDocument(BytesIO(data))
        return self._process_docx(doc, filename or "uploaded.docx", len(data), 
                                  "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    def _process_docx(self, doc, source_name: str, file_size: int, mime: str) -> ExtractedDocument:
        sections = []
        full_text_parts = []
        current_header = None
        current_text = []
        title = None
        
        # Try to get title from core properties
        if doc.core_properties.title:
            title = doc.core_properties.title
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            style = para.style.name if para.style else "Normal"
            is_heading = style.startswith('Heading') or text.upper() == text
            
            # Detect title from first heading if no core title
            if is_heading and not title and style.startswith('Heading 1'):
                title = text
            
            if is_heading:
                # Save previous section
                if current_text:
                    section_text = '\n'.join(current_text)
                    sections.append(ExtractedSection(
                        text=section_text,
                        section_header=current_header,
                        metadata={"style": style, "alignment": str(para.alignment)}
                    ))
                    full_text_parts.append(section_text)
                
                current_header = text
                current_text = []
            else:
                current_text.append(text)
        
        # Last section
        if current_text:
            section_text = '\n'.join(current_text)
            sections.append(ExtractedSection(
                text=section_text,
                section_header=current_header,
                metadata={"style": "Normal"}
            ))
            full_text_parts.append(section_text)
        
        # Extract tables as structured text
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(row_text))
            if rows:
                table_texts.append("\n".join(rows))
        
        return ExtractedDocument(
            source=source_name,
            doc_type=DocumentType.DOCX,
            title=title,
            content="\n\n".join(full_text_parts),
            sections=sections,
            metadata={
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "tables_text": table_texts if table_texts else None,
                "author": doc.core_properties.author,
                "created": str(doc.core_properties.created) if doc.core_properties.created else None,
            },
            file_size_bytes=file_size,
            mime_type=mime
        )